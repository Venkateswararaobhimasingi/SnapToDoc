import os
import tkinter as tk
from tkinter import filedialog, messagebox
import pyautogui
from docx import Document
from docx.shared import Inches
import winsound

# Global variables
document = None
document_path = ""
image_count = 0
start_x = start_y = end_x = end_y = 0
rect_id = None

def initialize_document():
    global document, document_path

    root = tk.Tk()
    root.withdraw()

    choice = messagebox.askquestion("Document Choice", "Do you want to use an existing document?")

    if choice == "yes":
        document_path = filedialog.askopenfilename(title="Select Word Document", filetypes=[("Word Documents", "*.docx")])
        if document_path:
            document = Document(document_path)
        else:
            messagebox.showerror("Error", "No document selected. Exiting.")
            exit()
    else:
        document_path = filedialog.asksaveasfilename(title="Create New Document", defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if document_path:
            document = Document()
            document.save(document_path)
        else:
            messagebox.showerror("Error", "No file name provided. Exiting.")
            exit()

def take_screenshot():
    global document, document_path, image_count

    root.withdraw()
    screenshot = pyautogui.screenshot()
    root.deiconify()

    save_screenshot(screenshot)

def take_custom_screenshot():
    global root, start_x, start_y, end_x, end_y, rect_id

    overlay = tk.Toplevel(root)
    overlay.attributes('-fullscreen', True)
    overlay.attributes('-alpha', 0.3)
    overlay.config(bg="black")

    canvas = tk.Canvas(overlay, cursor="cross", bg="black", highlightthickness=0)
    canvas.pack(fill=tk.BOTH, expand=True)

    def on_press(event):
        global start_x, start_y, rect_id
        start_x, start_y = event.x, event.y
        canvas.delete("rect")
        rect_id = canvas.create_rectangle(start_x, start_y, start_x, start_y, outline="red", width=2, fill="white", stipple="gray50")

    def on_drag(event):
        global end_x, end_y
        end_x, end_y = event.x, event.y
        canvas.coords(rect_id, start_x, start_y, end_x, end_y)

    def on_release(event):
        global start_x, start_y, end_x, end_y

        overlay.destroy()
        root.withdraw()

        x, y = min(start_x, end_x), min(start_y, end_y)
        width, height = abs(end_x - start_x), abs(end_y - start_y)

        if width > 0 and height > 0:
            screenshot = pyautogui.screenshot(region=(x, y, width, height))
            save_screenshot(screenshot)
        else:
            messagebox.showerror("Error", "Invalid selection. Try again.")

        root.deiconify()

    canvas.bind("<ButtonPress-1>", on_press)
    canvas.bind("<B1-Motion>", on_drag)
    canvas.bind("<ButtonRelease-1>", on_release)

def save_screenshot(screenshot):
    global document, document_path, image_count

    temp_image_path = f"temp_screenshot_{image_count}.png"
    screenshot.save(temp_image_path)

    document.add_picture(temp_image_path, width=Inches(5))
    image_count += 1
    document.save(document_path)
    os.remove(temp_image_path)
    winsound.PlaySound("SystemAsterisk", winsound.SND_ALIAS)

def add_text_from_gui():
    global document, document_path, text_input
    text = text_input.get("1.0", tk.END).strip()
    if text:
        document.add_paragraph(text)
        document.save(document_path)
        winsound.PlaySound("SystemAsterisk", winsound.SND_ALIAS)
        text_input.delete("1.0", tk.END)  # Clear text box after saving

def create_gui():
    global root, text_input
    root = tk.Tk()
    root.title("Screenshot to Word")
    root.geometry("200x200")
    root.configure(bg="#f0f0f0")
    root.attributes('-topmost', 1)
    root.overrideredirect(True)

    button_frame = tk.Frame(root, bg="#f0f0f0")
    button_frame.pack(pady=5, padx=10)

    tk.Button(button_frame, text="📸 Full", command=take_screenshot, font=("Arial", 10, "bold"), bg="#4CAF50", fg="white").grid(row=0, column=0, padx=5)
    tk.Button(button_frame, text="Custom", command=take_custom_screenshot, font=("Arial", 10, "bold"), bg="#FFA500", fg="white").grid(row=0, column=1, padx=3)

    text_input = tk.Text(root, height=2, width=25, font=("Arial", 10))
    text_input.pack(pady=5, padx=10, fill=tk.X)

    tk.Button(root, text="📝 Add Text", command=add_text_from_gui, font=("Arial", 10, "bold"), bg="#2196F3", fg="white").pack(pady=5, padx=10, fill=tk.X)
    tk.Button(root, text="❌ Exit", command=cancel_application, font=("Arial", 10, "bold"), bg="#f44336", fg="white").pack(pady=5, padx=10, fill=tk.X)

    def start_drag(event):
        root.x, root.y = event.x, event.y

    def drag_window(event):
        x, y = root.winfo_pointerx() - root.x, root.winfo_pointery() - root.y
        root.geometry(f"+{x}+{y}")

    root.bind('<Button-1>', start_drag)
    root.bind('<B1-Motion>', drag_window)
    root.mainloop()

def cancel_application():
    exit()
if __name__ == "__main__":
    initialize_document()
    create_gui()
